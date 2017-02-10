from PIL import Image
from os import listdir
from docx import Document
from docx.shared import Inches
import os

path="/home/ubuntu/Images"

document = Document()
p = document.add_paragraph()
for file in os.listdir(path):
   if file.endswith(".png"):
        print(file)
	
	r = p.add_run()
	r.add_picture(path+"/"+file)

document.save('demo.docx')

